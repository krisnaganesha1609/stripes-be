import io
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from docx import Document

class ExportService:
    @staticmethod
    def to_md(note):
        """Export note jadi Markdown string"""
        content = f"# {note['title']}\n\n"
        content += f"**Summary:**\n{note.get('summary', '')}\n\n"
        content += f"**Tags:** {', '.join(note.get('tags', []))}\n\n"
        content += f"**Content:**\n{note['content']}\n"
        return content

    @staticmethod
    def to_txt(note):
        """Export note jadi TXT string"""
        content = f"{note['title']}\n\n"
        content += f"Summary:\n{note.get('summary', '')}\n\n"
        content += f"Tags: {', '.join(note.get('tags', []))}\n\n"
        content += f"Content:\n{note['content']}\n"
        return content

    @staticmethod
    def to_pdf(note):
        """Export note jadi PDF (binary stream)"""
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=A4)
        textobject = c.beginText(40, 800)
        textobject.setFont("Helvetica", 12)

        lines = [
            f"Title: {note['title']}",
            "",
            f"Summary: {note.get('summary', '')}",
            "",
            f"Tags: {', '.join(note.get('tags', []))}",
            "",
            f"Content:",
            note['content']
        ]

        for line in lines:
            textobject.textLine(line)

        c.drawText(textobject)
        c.showPage()
        c.save()

        buffer.seek(0)
        return buffer

    @staticmethod
    def to_docx(note):
        """Export note jadi DOCX (binary stream)"""
        doc = Document()
        doc.add_heading(note['title'], 0)
        doc.add_paragraph(f"Summary: {note.get('summary', '')}")
        doc.add_paragraph(f"Tags: {', '.join(note.get('tags', []))}")
        doc.add_paragraph("Content:")
        doc.add_paragraph(note['content'])

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
